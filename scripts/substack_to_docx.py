#!/usr/bin/env python3
"""Convert a Surviving the Feds dispatch (markdown) into a Word document.

The output is just the article: an optional column graphic, the title, the
byline, and the body -- formatted so bold, italic, horizontal rules, and the
WARNING callout survive a paste into Substack's editor.

The column graphic is used automatically if it exists at one of the paths in
DEFAULT_GRAPHICS, or pass --graphic to point at a different file.

Usage:
    python3 scripts/substack_to_docx.py content/substack/<slug>.md
    python3 scripts/substack_to_docx.py content/substack/<slug>.md --graphic path/to/art.png
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Checked in order; first one that exists is used.
DEFAULT_GRAPHICS = [
    "assets/img/substack-column-header.png",
    "assets/img/substack-column-header.jpg",
    "assets/img/dispatches-header.png",
    "assets/img/dispatches-header.jpg",
]

WARNING_RED = RGBColor(0xB0, 0x00, 0x00)
TEXT_WIDTH = Inches(6.5)

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def horizontal_rule(doc):
    """Insert a thin horizontal rule as a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)


def add_rich_text(paragraph, text):
    """Render **bold** and *italic* markdown spans as real Word runs."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            paragraph.add_run(chunk[1:-1]).italic = True
        else:
            paragraph.add_run(chunk)


def parse(md_path):
    """Split the source file into title, subtitle, and body lines."""
    lines = Path(md_path).read_text().splitlines()
    title = subtitle = None
    body_start = 0

    for i, line in enumerate(lines):
        stripped = line.strip()
        if title is None and stripped.startswith("# "):
            title = stripped[2:].strip()
            continue
        if title and subtitle is None and stripped.startswith("*") and stripped.endswith("*"):
            subtitle = stripped.strip("*").strip()
            continue
        if stripped == "---":
            body_start = i + 1
            break

    if not title:
        sys.exit(f"error: no '# Title' heading found in {md_path}")

    return title, subtitle or "", lines[body_start:]


def find_graphic(explicit):
    if explicit:
        p = Path(explicit)
        if not p.exists():
            sys.exit(f"error: graphic not found at {p}")
        return p
    for candidate in DEFAULT_GRAPHICS:
        p = Path(candidate)
        if p.exists():
            return p
    return None


def build_header(doc, title, subtitle, graphic):
    if graphic:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(graphic), width=TEXT_WIDTH)

    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    if subtitle:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(12)

    byline = doc.add_paragraph()
    byline.add_run("By Bilal Khan")
    column = doc.add_paragraph()
    column.add_run("Surviving the Feds: Dispatches from the Inside").italic = True

    horizontal_rule(doc)


def build_body(doc, body_lines):
    for raw in body_lines:
        line = raw.strip()

        if not line:
            continue

        if line == "---":
            horizontal_rule(doc)
            continue

        if line.startswith("WARNING:"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = WARNING_RED
            continue

        # A lone bolded line is a section label, not prose.
        if line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
            continue

        add_rich_text(doc.add_paragraph(), line)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("--graphic", help="path to the column graphic image")
    args = ap.parse_args()

    md_path = Path(args.path)
    if not md_path.exists():
        sys.exit(f"error: {md_path} not found")

    title, subtitle, body_lines = parse(md_path)
    graphic = find_graphic(args.graphic)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    build_header(doc, title, subtitle, graphic)
    build_body(doc, body_lines)

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    print(f"wrote {out_path}" + (f" (graphic: {graphic})" if graphic else " (no graphic found)"))


if __name__ == "__main__":
    main()
